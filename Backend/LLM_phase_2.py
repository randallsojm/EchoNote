from langchain.text_splitter import RecursiveCharacterTextSplitter #Splits long texts into smaller chunks.
from langchain.chains.summarize import load_summarize_chain #Loads a summarization chain that helps analyze the text
from langchain_ollama import ChatOllama
from ollama import chat
from typing import List
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
import json

# Vector Store and retrievals
from langchain.chains import RetrievalQA
from pydantic import BaseModel, Field
import os
 
# Chat Prompt templates for dynamic values
# from langchain.prompts import PromptTemplate
from langchain.prompts.chat import (
    ChatPromptTemplate,
    SystemMessagePromptTemplate,
    HumanMessagePromptTemplate
)
from docx import Document
from langchain.schema import Document as LangChainDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Mm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH



def word_font(paragraph, run, Bold, alignment, underline, font_size):
    paragraph.alignment = alignment  # Set alignment before adding text
    run.font.name = "Times New Roman"  # Change font type
    run.font.size = Pt(font_size)  # Change font size
    run.bold = Bold            # Bold text     
    run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color (black)
    run.underline = underline

def set_cell_borders(cell, left_color, other_color):

    # Access the cell XML and ensure <w:tcPr> exists
    tc_pr = cell._element.get_or_add_tcPr()
    
    # Ensure <w:tcBorders> exists, if not, create it
    tc_borders = tc_pr.find(qn('w:tcBorders'))
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)

    # Function to create border elements
    def create_border(side, color):
        border = OxmlElement(side)
        border.set(qn('w:val'), 'single')  # Border style
        border.set(qn('w:sz'), '4')  # Border thickness
        border.set(qn('w:space'), '0')  # No space between the borders
        border.set(qn('w:color'), color)  # Border color
        return border

    # Create and append the borders
    tc_borders.append(create_border('w:top', other_color))  # Top border
    tc_borders.append(create_border('w:bottom', other_color))  # Bottom border
    tc_borders.append(create_border('w:right', other_color))  # Right border
    tc_borders.append(create_border('w:left', left_color))  # Left border
    

def add_contents_to_table(docx_path, topic_name, expanded_topic, row_idx, topic_item_no):
    # Open the existing docx file
    doc = Document(docx_path)
    
    # Check if the document contains any tables
    if doc.tables:
        # Select the first table in the document
        table = doc.tables[0]
        table.add_row()
        row_idx +=1
        print("row_idx", row_idx)
        cell = table.cell(row_idx, 0)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(f"ITEM {topic_item_no} - {topic_name}")
        word_font(paragraph, run, True, WD_ALIGN_PARAGRAPH.LEFT, False, 14)
        table.add_row()
        row_idx +=1
        cell_main_point = table.cell(row_idx, 0)
        paragraph = cell_main_point.paragraphs[0]
        run = paragraph.add_run(f"\n{expanded_topic}\n")
        word_font(paragraph, run, False, WD_ALIGN_PARAGRAPH.LEFT, False, 14)
        # Add who should take note in the second column
            
            
        doc.save(docx_path)
        return row_idx
    
    else:
        print("No tables found in the document.")
    

# Define a structured schema for extracted topics
class Topic(BaseModel):
    title: str = Field(..., description="Name of the topic")  # Adding description to the field
    description: str = Field(..., description="Brief description of the topic")  # Adding description to the field

class MeetingTopics(BaseModel):
    topics: List[Topic] = Field(..., description="List of extracted topics")  # Adding description to the list of topics

def LLM():
    # Create an instance of the OllamaLLM with LLaMA3 model
    # Initialize the LLM
    llm = ChatOllama(
        model="llama3:8b",
        temperature=0,
    )
    return llm

def text_splitter(all_chunks_output):
    print("all_chunks_output", all_chunks_output)
    for chunk in all_chunks_output:
        print(f"Processing chunk: {chunk}")
    docs = [LangChainDocument(page_content="".join(chunk)) for chunk in all_chunks_output]
    # print(f"You have {len(docs)} docs. First doc is {len(docs[0].page_content.split())} tokens")
    
    print("docs:", docs)
    return docs

def topic_extraction(llm, docs):
    # Define the system message prompt for extracting topics
    topic_extraction_template = """
    You are a helpful assistant that creates topics based on the context of a podcast transcript.
- Extract major ideas, themes, concepts, insights, comparisons, trends, viewpoints, and recurring points from the discussion.
- Each topic should be summarized with a brief, 1-sentence description of the idea or point.
- Use the exact language from the podcast for consistency.
- Do not include anything outside of the podcast discussion.
- Present each topic in the format: `Topic: Brief description`

Note: **If no clear topics are present, respond with "No Topics."**
    """
    system_message_prompt_map = SystemMessagePromptTemplate.from_template(topic_extraction_template)
    human_template = "Transcript: {text}"  # This will be dynamically replaced with each chunk of the transcript
    human_message_prompt_map = HumanMessagePromptTemplate.from_template(human_template)

    # Create the chat prompt for processing the transcript
    chat_prompt_map = ChatPromptTemplate.from_messages(messages=[system_message_prompt_map, human_message_prompt_map])

    # Define the system message prompt for deduplication and combining results
    combination_template = """
    You are an assistant that extracts and consolidates unique topics from a transcript.

- Given a list of topics with descriptions, extract unique topics and provide a brief, merged description for similar ones.
- **Remove duplicates** and avoid repeating the same ideas.
- Ensure each topic is **distinct** and clearly described.
- Return **only the topics and their descriptions**, without preamble or conclusions.

Note: **If no clear topics are present, respond with "No Topics."**

    """
    system_message_prompt_map_combined = SystemMessagePromptTemplate.from_template(combination_template)
    chat_prompt_combine = ChatPromptTemplate.from_messages(messages=[system_message_prompt_map_combined, human_message_prompt_map])

    # Create summarization chain
    chain = load_summarize_chain(
        llm,
        chain_type="map_reduce",
        map_prompt=chat_prompt_map,
        combine_prompt=chat_prompt_combine,
    )

    # Perform topic extraction
    topics_found = chain.invoke({"input_documents": docs})
    topics_found = topics_found.get("output_text", "")
    print("topics:", topics_found )
    return topics_found

def jsonify(topics_found, model_class, schema):
    prompt =  f"""
    Your task is to **parse** and **format** this content into the appropriate structure as defined by the provided schema, without any additional transformations or summarization.
    Please **retain** the original content as it is, without summarizing or altering any of the text. 

    Content:
    {topics_found}
    """
    json_response = chat(
    messages=[
        {
        'role': 'user',
        'content': prompt,
        }
    ],
    model='llama3:8b',
    format=model_class.model_json_schema(),
    )
    # print(f"Response content: {json_response.message.content}")

    topics_list = model_class.model_validate_json(json_response.message.content)
    #topics_list is an instance of MeetingTopics
    #topics_list itself isn't a list—it's an object with a .topics attribute that holds the list.
    topics_dict = topics_list.model_dump()
    topics_dict = topics_dict.get(schema, [])
    file_path = "topics_list_test.json"
    # Open the file in write mode and write the list of dictionaries as JSON
    with open(file_path, "w") as file:
        json.dump(topics_dict, file, indent=4)  
        # Write the list in a human-readable format
        #indent=4 means each nested element (like dictionaries within a list) will be indented by 4 spaces for better readability
    return topics_dict


def RAG_for_summary(llm, docs, topics_dict, row_idx, docx_path):
    embedding_model = "all-MiniLM-L6-v2"
    model_kwargs = {"device": "cuda"}
    embeddings = HuggingFaceEmbeddings(
    model_name=embedding_model,
    model_kwargs=model_kwargs)

    # Store document embeddings in FAISS
    vectorstore = FAISS.from_documents(docs, embeddings)


    # Step 2: Expand on the topics you found
    # In order to expand on the topics we found, we are going to use vectorstore for context-based retrieval



    # Define system template for the retrieval-based summary
    summary_template = """
You will be given a podcast transcript containing multiple topics. Your task is to summarize key points for the topic the user chooses.

For each key point, provide a single, detailed paragraph (maximum 5 sentences) covering the following elements:

- Introduce the topic briefly based on the discussion.
- Identify challenges or problems raised, specifying who is facing them.
- Summarize the actions each speaker is taking and the progress of any related projects or initiatives.
- Outline any proposed ideas or solutions and the speaker suggesting them.
- Note any disagreements and the speakers involved.
- List follow-up actions and assign responsibility where applicable.

If any of the above details are missing, omit them rather than making assumptions.
Avoid preambles, introductory statements, and concluding remarks. Exclude irrelevant tangents or off-topic discussions.
Follow the format below:
### output format
```
1. <key point 1><One detailed paragraph about key point 1>

2. <key point 2><One detailed paragraph about key point 2>
```
-----------------
{context}
"""

    formatting_prompt = """
Format the provided text for improved readability and presentation without changing the content

Follow these guidelines for formatting:

1. Bold Key Points:
- Bold the key points to make them stand out clearly.
- After each bolded key point, provide a detailed paragraph explaining the point concisely.

2. Clear Section Breaks:
- Separate each key point with clear breaks (i.e., extra space between sections).
- This ensures each point is easily distinguishable, helping readers navigate the content.

3.Use Bullet Points for Related Ideas:
- If there are multiple related actions, ideas, or points, break them down using bullet points for clarity.

4. Use Numbered Lists for Sequential Steps:
- If presenting a sequence or order of steps, use numbered lists to make the progression clear.
"""

    messages_RAG = [
        SystemMessagePromptTemplate.from_template(summary_template),
        HumanMessagePromptTemplate.from_template("{question}"),
    ]

    formatting_prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                formatting_prompt,
            ),
            ("human", "{input}"),
        ]
        )
    formatting_chain = formatting_prompt | llm
    # Invoke the LLM to extract decisions and actions
   


    # This will pull the two messages together and get them ready to be sent to the LLM through the retriever
    CHAT_PROMPT_RAG = ChatPromptTemplate.from_messages(messages_RAG)
    # Perform the retrieval-based summary using GPT-4
    qa = RetrievalQA.from_chain_type(
        llm,
        chain_type="stuff",
        retriever=vectorstore.as_retriever(k=4),
        chain_type_kwargs={'prompt': CHAT_PROMPT_RAG}
    )


    if not topics_dict:
        print("No topics found.")
        doc = Document(docx_path)
        if doc.tables:
        # Select the first table in the document
            table = doc.tables[0]
            table.add_row()
            row_idx +=1
            print("row_idx", row_idx)
            cell = table.cell(row_idx, 0)
            paragraph = cell.paragraphs[0]
            expanded_topic_str = "Transcript is too short. There are no topics to analyse."
            run = paragraph.add_run(expanded_topic_str)
            word_font(paragraph, run, False, WD_ALIGN_PARAGRAPH.LEFT, False, 14)
            doc.save(docx_path)
            row_idx = None
    else:
        # Process each topic
        expanded_topic_list = []
        for topic_item_no, topic in enumerate(topics_dict, start = 1):  # Limiting to first 3 topics for conciseness
            topic_name = topic.get("title", "No title")
            topic_description = topic.get("description", "No description")
            print(topic_name)

            # Construct query for expanding the topic
            query = f"""
                {topic_name}: {topic_description}
            """


            # Retrieve expanded topic information
            expanded_topic = qa.invoke({"query" : query})
            expanded_topic = expanded_topic.get("result", '')
            formatted_summary = formatting_chain.invoke({"input": expanded_topic})  # Uses the same docs from transcript processing
            formatted_summary = formatted_summary.content
            expanded_topic_list.append(formatted_summary)
            print("formatted_summary", formatted_summary)
            row_idx = add_contents_to_table(docx_path, topic_name, expanded_topic, row_idx, topic_item_no)
        expanded_topic_str = '\n'.join(expanded_topic_list)
    return row_idx, expanded_topic_str

            
            
            
# expanded_topic_list: [[{'main_point': 'Traditional firewalls and intrusion detection systems are becoming obsolete against deep learning based cyber threats.', 'who_should_take_note': '*ALL*'}, {'main_point': 'We need a paradigm shift in how we define critical infrastructure to keep up with AI driven attacks.', 'who_should_take_note': '*ALL*'}, {'main_point': 'The problem is policy resistance, as some officials still view AI in cybersecurity as a black box risk.', 'who_should_take_note': '*ALL*'}, {'main_point': 'A risk assessment report highlighting recent AI driven breaches should be prepared for the next meeting with the cyber security task force.', 'who_should_take_note': '*ALL*'}, {'main_point': 'Clearer rules are needed to regulate the integration of AI in combat systems, including fully autonomous drones and AI assist targeting systems.', 'who_should_take_note': '*ALL*'}, {'main_point': 'A policy framework that mandates human-in-the-loop oversight is necessary for AI system decisions.', 'who_should_take_note': '*ALL*'}, 
# Topics dictionary: {'expanded_topics': [{'main_point': 'Regional Stability', 'who_should_take_note': 'ALL;Helin;Randall'}]}
# Topics dictionary: {'expanded_topics': [{'main_point': 'Cyber Security', 'who_should_take_note': 'ALL'}, {'main_point': 'AI in Warfare', 'who_should_take_note': 'ALL'}, {'main_point': 'Regional Stability', 'who_should_take_note': 'Randall, Helin'}, {'main_point': 'Hybrid Warfare', 'who_should_take_note': 'Randall, Helin'}]}


        # Load the text splitter for chunking the transcript into manageable parts
def speech_to_text(llm, expanded_topic_str, row_idx):
    if row_idx is not None:
        speech_to_text_template = """
        You will be given a summary of a podcast transcript, consisting of topics with their descriptions and follow-up actions.

        You are an AI assistant that converts the summarized content into natural and clear speech for listening.

        Your task is to:  
        - **Start with a simple and professional introduction**, such as:  
        *"Welcome to today’s briefing. In this episode, we will cover critical updates and key discussions from our latest meeting. Let’s begin by reviewing the first topic."*  
        - **Convert each topic and its description into a natural, conversational speech** that flows smoothly.  
        - Start with *"First, we have [Topic 1]..."* for the first topic and use *"Next, we have..."* for subsequent topics.  
        - **After covering all topics, transition into follow-up actions** with a phrase like:  
        *"Now, let’s go over the key actions from this meeting."*  
        - **Convert each action into a clear and engaging spoken format** while maintaining conciseness.  
        - Use phrases like *"[Speaker] will be responsible for..."* or *"[Speaker] will follow up by..."* to ensure clarity.  
        - **End with a concluding sentence**, such as:  
        *"That wraps up today’s briefing. Stay tuned for our next update."*  
        - **Do not change the meaning of the text.**  
        - **Only return the output meant for the podcast.** Do not include preamble, explanations, or any content irrelevant to the spoken output.  

    """
        prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                speech_to_text_template,
            ),
            ("human", "{input}"),
        ]
        )
        Analysis_speech_chain = prompt | llm
        # Invoke the LLM to extract decisions and actions
        Analysis_speech = Analysis_speech_chain.invoke({"input": expanded_topic_str})  # Uses the same docs from transcript processing
        # print("Analysis_speech:", Analysis_speech)
        Analysis_speech = Analysis_speech.content
    else:
        Analysis_speech = expanded_topic_str
    with open("processed/podcast.txt", "w")as f:
        f.write(Analysis_speech)
    return Analysis_speech


def format_table_of_contents(docx_path):
    doc = Document(docx_path)
    # Check if the document contains any tables
    if doc.tables:
        # Select the first table in the document
        table = doc.tables[0]
        for row in table.rows:
        # Access the second column (index 1)
            second_column_cell = row.cells[1]
            set_cell_borders(second_column_cell, '000000', 'FFFFFF') # Apply the custom border settings to the second cell
            first_column_cell = row.cells[0]
            set_cell_borders(first_column_cell, 'FFFFFF', 'FFFFFF')
            # Save the document
        doc.save(docx_path)

class Speaker(BaseModel):
    label: str = Field(..., description="Speaker label (e.g., 'Speaker A')")
    name: str = Field(None, description="Extracted speaker name, or 'Unknown' if not introduced")

class SpeakerMapping(BaseModel):
    speakers: List[Speaker] = Field(..., description="List of speakers with their labels and names")


def speaker_identification(llm, transcript_path):
    system_template = """

You are given a meeting transcript with labeled speakers (e.g., Speaker A, Speaker B, Speaker C). 

**Task:** Extract and match speaker labels to names based strictly on self-introductions and contextual mentions.   
**Extraction Rules:**
1. Use self-introductions first (e.g., "Hi, I am John").
2. If absent, use contextual mentions (e.g., "Welcome, Sarah!").
3. Do not infer or assume names—only extract explicitly stated ones.
4. If no name is found, label as "Unknown."

Output the result in this format:  
   ```
   Speaker A: [Name]  
   Speaker B: [Name]  
   Speaker C: [Name]  
   ```
""" 
    with open(transcript_path, "r") as f:
        transcript = f.read()
    prompt = ChatPromptTemplate.from_messages(
    [
        (
            "system",
            system_template,
        ),
        ("human", "{input}"),
    ]
    )
    speaker_identification_chain = prompt | llm
    # Invoke the LLM to extract decisions and actions
    speaker_identification = speaker_identification_chain.invoke({"input": transcript})  # Uses the same docs from transcript processing
    # print("Analysis_speech:", Analysis_speech)
    speaker_identification = speaker_identification.content
    print(speaker_identification)
    return speaker_identification




# docx_path = 'Successful_trial_1.docx'
# transcript_path = "Sample_Transcript.txt"

row_idx = 0
def main_LLM(meeting_starttime, meeting_endtime, transcript_path, row_idx, docx_path, all_chunks_output):
    llm = LLM()
    speaker_mapping = speaker_identification(llm, transcript_path)
    speaker_mapping_list = jsonify(speaker_mapping, SpeakerMapping, 'speakers')
    print("speaker_mapping_list", speaker_mapping_list)
    speaker_dict = {speaker['label'].replace('Speaker ', ''): speaker['name'] for speaker in speaker_mapping_list}
    all_chunks_output_replaced = []
    all_chunk_output_presentable = []
    for chunk in all_chunks_output:
        chunk = "".join(chunk)
        print("chunk", chunk)
        for label, name in speaker_dict.items():
            print("label", label)
            if f" | {label}" in chunk:
                chunk.replace(f" | {label}", f" | {name}")
            all_chunks_output_replaced.append([chunk])
            all_chunk_output_presentable.append(chunk)


    all_chunk_output_presentable = "".join(all_chunk_output_presentable)
    documents_folder = 'processed'
    if not os.path.exists(documents_folder):
        os.makedirs(documents_folder)
    formatted_transcript = f"{documents_folder}/Transcript.txt"
    with open(formatted_transcript, 'w') as f:
        f.write(all_chunk_output_presentable)

    speakers = ", ".join([speaker['name'] if speaker.get('name') not in [None, 'Unknown'] else speaker['label'] for speaker in speaker_mapping_list])
# If name is provided (not None), it will output the actual name.
# If name is None or 'unknown', it will output the speaker label
    docs = text_splitter(all_chunks_output_replaced)
    topics_found = topic_extraction(llm, docs)
    topics_dict = jsonify(topics_found, MeetingTopics, 'topics')
    
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("EchoNote\n\n")
    word_font(paragraph, run, True, WD_ALIGN_PARAGRAPH.CENTER, True, 20)
    if meeting_starttime is not None and meeting_endtime is not None:
        run = paragraph.add_run("Start of Meeting: ")
        word_font(paragraph, run, False, WD_ALIGN_PARAGRAPH.LEFT, False, 14)
        run = paragraph.add_run(f"{meeting_starttime} hrs\n")
        word_font(paragraph, run, True, WD_ALIGN_PARAGRAPH.LEFT, False, 14)
        run = paragraph.add_run("End of Meeting: ")
        word_font(paragraph, run, False, WD_ALIGN_PARAGRAPH.LEFT, False, 14)
        run = paragraph.add_run(f"{meeting_endtime} hrs\n")
        word_font(paragraph, run, True, WD_ALIGN_PARAGRAPH.LEFT, False, 14)
    run = paragraph.add_run("Speakers: ")
    word_font(paragraph, run, False, WD_ALIGN_PARAGRAPH.LEFT, False, 14)
    run = paragraph.add_run(f"{speakers}\n")
    word_font(paragraph, run, True, WD_ALIGN_PARAGRAPH.LEFT, False, 14)
    run = paragraph.add_run(f"Meeting Topics:\n")
    word_font(paragraph, run, False, WD_ALIGN_PARAGRAPH.LEFT, False, 14)
    for idx, topic in enumerate(topics_dict):
        topic_name = topic.get('title', 'No topic name')
        run = paragraph.add_run(f"{idx+1}.  {topic_name}\n")
        word_font(paragraph, run, True, WD_ALIGN_PARAGRAPH.LEFT, False, 14)
        topic_description = topic.get('description', 'No description')
        run = paragraph.add_run(f"{topic_description}\n\n")
        word_font(paragraph, run, False, WD_ALIGN_PARAGRAPH.LEFT, False, 14)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("DESCRIPTION")
    word_font(paragraph, run, True, WD_ALIGN_PARAGRAPH.JUSTIFY, True, 14)
    doc.save(docx_path)

    # expanded_topic_list = RAG_for_expanded_topic(llm, docs, topics_dict)
    row_idx, expanded_topic_str = RAG_for_summary(llm, docs, topics_dict, row_idx, docx_path)
    podcast_text = speech_to_text(llm, expanded_topic_str, row_idx)

    print("LLM_phase_2 script run completed")
    return podcast_text

transcript_path = "Transcript.txt"
docx_path = 'Minutes_Echonote.docx'
all_chunks_output =  [["0.00 --> 0.10 | A\n Alright, let's start with cyber security.  Probably the most urgent topic right now.  We have seen a 40% increase in cyber incursions just in the past year.\n\n"], ["0.10 --> 0.20 | B\n Right, and what's most concerning is the shift towards AI driven attacks.  These are just isolated features anymore.  They are automated, persistent, and adaptive.\n\n"]]
main_LLM('1200', '1400', transcript_path, row_idx, docx_path, all_chunks_output)

# docx_path = 'processed/Minutes_Echonote.docx'
# main_LLM('1200', '1400', 'formatted_transcript.txt', row_idx, docx_path, all_chunks_output)
# expanded_topic_str = '\n\n'.join(expanded_topic_list)
# actions_found = action_extraction(llm, docs)
# Analysis_speech = speech_to_text(llm, expanded_topic_str)


# schema = {
    # "properties": {
    #     "topic_name": {
    #         "type": "string",
    #         "description": "The title of the topic listed"
    #     },
    #     "description": {
    #         "type": "string",
    #         "description": "The description of the topic listed"
    #     },
    # },
    # "required": ["topic_name", "description"],
    # }

    # # Create prompt to format summary as JSON according to schema
    # prompt = f"""
    # You will be given a summary of a transcript, consisting of topic names and their descriptions. 
    # From the summary, extract out the topic names and descriptions and give an output.
    # The output should be formatted as a JSON instance that conforms to the following JSON schema:
    # {json.dumps(schema, indent=2)}
    # Summary: {topics_found}

    # For example, the summary below:
    # "Cyber Security - The most urgent topic right now, with a 40% increase in cyber incursions just in the past year.
    # AI Driven Attacks - Automated, persistent, and adaptive attacks that traditional firewalls and intrusion detection systems can't handle."

    # should output:

    # [
    #     {{
    #         "topic_name": "Cyber Security",
    #         "description": "The most urgent topic right now, with a 40% increase in cyber incursions just in the past year."
    #     }},
    #     {{
    #         "topic_name": "AI Driven Attacks",
    #         "description": "Automated, persistent, and adaptive attacks that traditional firewalls and intrusion detection systems can't handle."
    #     }}
    # ]

    # """



def action_extraction(llm, docs):
    action_prompt_template = """
    You are an AI assistant that extracts actions to be taken from a meeting transcript, which contains speaker labels and timestamps.
    Your task is to:
    - Identify all actions to be done, including tasks, next steps, or follow-up activities discussed during the meeting.
    - Present the output **strictly** in a numbered list format as shown below.
    - Each action should be described clearly and concisely.
    - Specify who is responsible for each action  by incorporating the speaker labels based on speaker labels and timestamps. 
    - If no actions were mentioned, respond with 'No actions to be taken.'


    ### What constitutes an action:
    - Tasks that need to be completed.
    - Steps or procedures that need to be followed.
    - Assignments or delegations of responsibility.
    - Follow-up activities or reminders.
    - Any other explicit or implied activity that needs to be executed based on the meeting.


    ### Format:
    Strictly return in this format:
    ```
    Follow-up actions include:

    1. "Prepare the report by next Tuesday" - [Speaker 1]
    2. "Send the updated document to the team" - [Speaker 2]
    3. "Follow up with the client next week" - [Speaker 1]
    ```

    **Important**: The actions must be **clearly numbered**. Each action should be **on its own line**. 
    **Important**: The actions must be written **accompanied with the speaker responsible for the task**.
    """

    # Create a prompt template for decision and action extraction
    action_prompt = ChatPromptTemplate.from_messages([
        SystemMessagePromptTemplate.from_template(action_prompt_template),
        HumanMessagePromptTemplate.from_template("Transcript: {text}")
    ])


    # Process the transcript to extract decisions and actions
    action_chain = load_summarize_chain(
        llm,
        chain_type="map_reduce",  # Use map_reduce instead of stuff
        map_prompt=action_prompt
    )

    # map_reduce indicates that the input (meeting transcript) will be processed in chunks (mapped), and then those chunks will be combined (reduced) to generate the final output.
    # The map_prompt is set to the action_prompt, which means each chunk of the transcript will be processed using the action extraction template.

    # Invoke the LLM to extract decisions and actions
    actions_found = action_chain.invoke({"input_documents": docs})  # Uses the same docs from transcript processing
    # print(actions_found)
    actions_found = actions_found.get('output_text', [])
    return actions_found